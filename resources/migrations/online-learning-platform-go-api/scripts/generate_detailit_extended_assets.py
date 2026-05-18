from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "resources" / "course_files" / "detailit"
FONT_REGULAR = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size=size)


def wrap(draw, text, fnt, max_width):
    lines = []
    for paragraph in text.split("\n"):
        current = ""
        for word in paragraph.split():
            candidate = f"{current} {word}".strip()
            if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        lines.append("")
    return lines[:-1]


def write_block(draw, x, y, text, fnt, fill, max_width, gap=8):
    for line in wrap(draw, text, fnt, max_width):
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + gap
    return y


def make_diagram(filename, title, subtitle, blocks, accent):
    img = Image.new("RGB", (1600, 1000), "#f4f6f8")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((62, 54, 1538, 946), radius=28, fill="#ffffff", outline="#d6dde7", width=3)
    d.rectangle((62, 54, 1538, 205), fill=accent)
    d.text((112, 88), title, font=font(50, True), fill="#ffffff")
    d.text((114, 155), subtitle, font=font(26), fill="#eef2ff")

    positions = [(112, 270), (830, 270), (112, 610), (830, 610)]
    for index, ((heading, body), (x, y)) in enumerate(zip(blocks, positions), start=1):
        d.rounded_rectangle((x, y, x + 645, y + 260), radius=22, fill="#f9fafb", outline="#d9e1ec", width=2)
        d.rounded_rectangle((x + 26, y + 28, x + 105, y + 107), radius=18, fill=accent)
        d.text((x + 65, y + 67), str(index), font=font(34, True), fill="#ffffff", anchor="mm")
        d.text((x + 132, y + 34), heading, font=font(31, True), fill="#1f2937")
        write_block(d, x + 132, y + 88, body, font(23), "#5f6b7a", 455)

    img.save(OUT / filename, quality=95)


def make_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Heading 1", 17, RGBColor(21, 94, 117)),
        ("Heading 2", 13, RGBColor(31, 41, 55)),
    ]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = color

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ООО "ДетаЛит"\nРегламент обучения подразделений по ОКВЭД 29.31')
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(21, 94, 117)

    doc.add_paragraph(
        "Документ описывает учебную сетку для сотрудников, которые работают с электрическим "
        "и электронным оборудованием для автотранспортных средств: жгутами, разъемами, "
        "электронными модулями, испытаниями, прослеживаемостью и входным контролем."
    )

    doc.add_heading("1. Роли и обязательные курсы", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for idx, heading in enumerate(["Роль", "Подразделение", "Базовый курс", "Практическая проверка"]):
        cell = table.rows[0].cells[idx]
        cell.text = heading
        cell.paragraphs[0].runs[0].bold = True
    for row in [
        ["Оператор", "Производственный участок", "Автоэлектроника и маршрут изделия", "Сборка узла по карте"],
        ["Сборщик", "Жгуты и разъемы", "Жгуты, маркировка и обжим", "Контроль обжима и пиновки"],
        ["Инженер", "Технологический отдел", "ESD, пайка и монтаж", "Разбор дефекта по 5 Why"],
        ["Контролер", "ОТК и лаборатория", "Испытания и трассируемость", "Оформление несоответствия"],
    ]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    doc.add_heading("2. Матрица допуска", level=1)
    for item in [
        "Сотрудник допускается к операции только после вводного курса, практической проверки и назначения мастером.",
        "Операции с ESD-риском требуют отдельного подтверждения знаний по защите компонентов.",
        "Работа с измерительным оборудованием требует знания методики, допуска и правил записи результата.",
        "Каждая партия должна быть связана с маршрутной картой, номером изделия, сменой и ответственным.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Производственные схемы", level=1)
    doc.add_picture(str(OUT / "electronics-architecture.png"), width=Inches(6.2))
    doc.add_picture(str(OUT / "harness-assembly.png"), width=Inches(6.2))

    doc.save(OUT / "detailit-department-training-regulation.docx")


def make_pdf():
    pdfmetrics.registerFont(TTFont("Arial", FONT_REGULAR))
    pdfmetrics.registerFont(TTFont("Arial-Bold", FONT_BOLD))

    path = OUT / "detailit-esd-soldering-standard.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontName="Arial-Bold", fontSize=20, leading=24, textColor=colors.HexColor("#155e75"), alignment=TA_CENTER)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName="Arial-Bold", fontSize=13, leading=16, textColor=colors.HexColor("#1f2937"), spaceBefore=10)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName="Arial", fontSize=10, leading=13)

    story = [
        Paragraph('ООО "ДетаЛит"', h1),
        Paragraph("Стандарт ESD, пайки и монтажа электронных компонентов", h1),
        Spacer(1, 8),
        Paragraph("Цель", h2),
        Paragraph("Снизить риск скрытых отказов электронного оборудования для автотранспортных средств за счет контроля ESD, подготовки рабочего места, качества пайки и прослеживаемости операций.", body),
        Paragraph("Контроль рабочего места", h2),
    ]

    data = [
        ["Контроль", "Норма", "Действие при отклонении"],
        ["ESD-браслет", "Проверка перед сменой", "Остановить монтаж, заменить браслет"],
        ["Коврик и тара", "Антистатическое исполнение", "Убрать обычную тару с поста"],
        ["Паяльная станция", "Температура по техкарте", "Откалибровать или передать мастеру"],
        ["Плата", "Без влаги и загрязнений", "Очистить, просушить, оформить замечание"],
    ]
    table = Table(data, colWidths=[45 * mm, 58 * mm, 76 * mm])
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "Arial"),
        ("FONTNAME", (0, 0), (-1, 0), "Arial-Bold"),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#155e75")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(table)
    story.extend([
        Paragraph("Критерии качества пайки", h2),
        Paragraph("Соединение должно быть блестящим или равномерно матовым по типу припоя, без мостиков, шариков, трещин, перегрева площадки и непропая. Спорная пайка не передается дальше без повторного осмотра.", body),
        Paragraph("Прослеживаемость", h2),
        Paragraph("В записи операции указываются номер партии, изделие, исполнитель, смена, результат контроля, оборудование и версия технологической карты.", body),
    ])
    doc.build(story)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    make_diagram(
        "electronics-architecture.png",
        "Электронное оборудование автомобиля",
        "Учебная схема для подразделений ООО \"ДетаЛит\"",
        [
            ("Датчики и входы", "Сигналы температуры, положения, напряжения и состояния цепей должны быть стабильными и проверяемыми."),
            ("ЭБУ и логика", "Блок управления принимает сигналы, выполняет алгоритм и передает команды исполнительным цепям."),
            ("Жгуты и разъемы", "Качество обжима, фиксации и маркировки влияет на надежность автомобиля сильнее, чем кажется."),
            ("Испытания", "Изделие проходит функциональную проверку, электрические измерения и запись результата в маршрут."),
        ],
        "#155e75",
    )
    make_diagram(
        "harness-assembly.png",
        "Сборка жгутов и разъемов",
        "Маркировка, обжим, пиновка и контроль",
        [
            ("Подготовка провода", "Проверить сечение, цвет, длину, маркировку и состояние изоляции до обжима."),
            ("Обжим контакта", "Контакт фиксируется без повреждения жилы; усилие и геометрия сверяются с картой."),
            ("Пиновка разъема", "Каждый провод ставится в свой контактный номер, ошибка пина блокирует выпуск изделия."),
            ("Финальный контроль", "Тест цепи, натяжение, внешний осмотр и запись результата обязательны для партии."),
        ],
        "#0f766e",
    )
    make_diagram(
        "esd-workstation.png",
        "ESD-пост монтажа",
        "Защита электронных компонентов от скрытых повреждений",
        [
            ("Браслет", "Проверяется перед сменой и после перерыва; без проверки монтаж не начинается."),
            ("Коврик", "Рабочая поверхность должна быть антистатической, чистой и подключенной к заземлению."),
            ("Тара", "Компоненты хранятся только в ESD-таре, обычные пакеты и коробки убираются с поста."),
            ("Запись", "Операция фиксируется с партией, исполнителем, оборудованием и результатом контроля."),
        ],
        "#7c3aed",
    )
    make_diagram(
        "testing-traceability.png",
        "Испытания и трассируемость",
        "Как доказать качество изделия по данным партии",
        [
            ("Маршрут", "Номер партии связывает изделие с материалами, операциями, сотрудниками и сменой."),
            ("Измерение", "Результат должен иметь методику, прибор, допуск и подпись ответственного."),
            ("Несоответствие", "Спорное изделие изолируется, причина и действие фиксируются до решения ОТК."),
            ("Архив", "Документы партии позволяют быстро восстановить историю при рекламации заказчика."),
        ],
        "#b45309",
    )
    make_docx()
    make_pdf()


if __name__ == "__main__":
    main()
