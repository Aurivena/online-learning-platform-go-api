from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "resources" / "course_files" / "detailit"
FONT_REGULAR = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def font(size, bold=False):
    path = FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(path, size=size)


def multiline(draw, xy, text, fnt, fill, width, line_gap=8):
    x, y = xy
    for para in text.split("\n"):
        lines = []
        for raw in para.split(" "):
            current = (lines[-1] + " " + raw).strip() if lines else raw
            if draw.textbbox((0, 0), current, font=fnt)[2] <= width:
                if lines:
                    lines[-1] = current
                else:
                    lines.append(current)
            else:
                lines.append(raw)
        for line in lines or [""]:
            draw.text((x, y), line, font=fnt, fill=fill)
            y += fnt.size + line_gap
        y += line_gap
    return y


def make_png(filename, title, subtitle, steps, accent):
    img = Image.new("RGB", (1600, 1000), "#f6f7f9")
    d = ImageDraw.Draw(img)
    navy = "#1f2937"
    muted = "#6b7280"
    d.rounded_rectangle((70, 60, 1530, 940), radius=28, fill="#ffffff", outline="#d8dee8", width=3)
    d.rectangle((70, 60, 1530, 210), fill=accent)
    d.text((120, 96), title, font=font(54, True), fill="#ffffff")
    d.text((122, 165), subtitle, font=font(28), fill="#eef2ff")
    y = 275
    for idx, (name, body) in enumerate(steps, start=1):
        x = 125 + ((idx - 1) % 2) * 700
        if idx == 3:
            y = 620
        box = (x, y, x + 610, y + 245)
        d.rounded_rectangle(box, radius=22, fill="#f9fafb", outline="#d9e0ea", width=2)
        d.ellipse((x + 28, y + 28, x + 96, y + 96), fill=accent)
        d.text((x + 52, y + 44), str(idx), font=font(30, True), fill="#ffffff", anchor="mm")
        d.text((x + 122, y + 32), name, font=font(32, True), fill=navy)
        multiline(d, (x + 122, y + 84), body, font(23), muted, 435, line_gap=7)
    img.save(OUT / filename, quality=95)


def make_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ООО "ДетаЛит"\nКарта ежедневного инструктажа сотрудника')
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph(
        "Документ предназначен для вводного и повторного обучения сотрудников литейного, "
        "механического и контрольного участков. Используйте его как чек-лист перед началом смены."
    )

    doc.add_heading("1. Перед началом смены", level=1)
    for item in [
        "Проверить спецодежду, защитные очки, перчатки и обувь.",
        "Осмотреть рабочую зону: проходы свободны, аварийные кнопки доступны.",
        "Проверить исправность вытяжки, ограждений и маркировки опасных зон.",
        "Получить задание смены и уточнить критерии качества партии.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Контрольные точки процесса", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Этап", "Что проверяем", "Норма", "Действие при отклонении"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    rows = [
        ["Плавка", "Температура и состав", "По техкарте", "Сообщить мастеру, остановить заливку"],
        ["Форма", "Чистота, смазка, фиксация", "Без сколов и перекосов", "Заменить форму или очистить"],
        ["Заливка", "Скорость и стабильность", "Без рывков", "Зафиксировать причину, отделить партию"],
        ["Охлаждение", "Время выдержки", "По регламенту", "Не вскрывать форму досрочно"],
        ["Контроль", "Геометрия и дефекты", "В пределах допуска", "Оформить несоответствие"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    doc.add_heading("3. Лист самопроверки", level=1)
    for item in [
        "Я понимаю маршрут эвакуации и место сбора.",
        "Я знаю, где находятся огнетушитель, аптечка и аварийная кнопка.",
        "Я умею отличить допустимый дефект от брака по карте качества.",
        "Я знаю, кому сообщать об отклонениях и как фиксировать событие.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("4. Иллюстрация процесса", level=1)
    doc.add_picture(str(OUT / "process-flow.png"), width=Inches(6.2))

    doc.save(OUT / "detailit-instruction-card.docx")


def make_pdf():
    pdfmetrics.registerFont(TTFont("Arial", FONT_REGULAR))
    pdfmetrics.registerFont(TTFont("Arial-Bold", FONT_BOLD))

    path = OUT / "detailit-quality-standard.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=15 * mm,
        leftMargin=15 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "h1",
        parent=styles["Heading1"],
        fontName="Arial-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#1f4e79"),
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "h2",
        parent=styles["Heading2"],
        fontName="Arial-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#263238"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName="Arial", fontSize=10, leading=13)
    small = ParagraphStyle("small", parent=body, fontSize=9, leading=12)

    story = [
        Paragraph('ООО "ДетаЛит"', h1),
        Paragraph("Стандарт визуального и измерительного контроля отливок", h1),
        Paragraph(
            "Материал описывает минимальный порядок контроля партии после охлаждения, очистки и первичной механической обработки.",
            body,
        ),
        Spacer(1, 6),
        Paragraph("Классы дефектов", h2),
    ]
    data = [
        ["Класс", "Признак", "Риск", "Решение"],
        ["A", "Трещины, сквозные раковины", "Высокий", "Брак, изоляция партии"],
        ["B", "Отклонение размера за пределами допуска", "Средний", "Повторное измерение, решение ОТК"],
        ["C", "Поверхностные следы без влияния на геометрию", "Низкий", "Зачистка и повторный осмотр"],
    ]
    table = Table(data, colWidths=[18 * mm, 63 * mm, 28 * mm, 64 * mm])
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), "Arial"),
                ("FONTNAME", (0, 0), (-1, 0), "Arial-Bold"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#b7c1cc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f7fb")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(table)
    story.extend(
        [
            Paragraph("Порядок контроля партии", h2),
            Paragraph(
                "1. Сверить номер партии, материал, маршрутную карту и актуальную версию чертежа. "
                "2. Выполнить визуальный контроль на освещённом посту. 3. Измерить критические размеры "
                "калибрами или штангенинструментом. 4. Зафиксировать результат в журнале ОТК. "
                "5. При отклонении отделить изделие и уведомить мастера смены.",
                body,
            ),
            Paragraph("Минимальные критерии приемки", h2),
        ]
    )
    for item in [
        "На поверхности нет трещин, сквозных пор, незаполнений и следов перегрева.",
        "Посадочные и базовые размеры находятся в пределах допуска чертежа.",
        "Маркировка партии читается, изделие можно однозначно связать с маршрутной картой.",
        "Решение по спорным дефектам принимает ОТК вместе с мастером участка.",
    ]:
        story.append(Paragraph("• " + item, small))
        story.append(Spacer(1, 2))
    doc.build(story)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    make_png(
        "safety-zone.png",
        "Безопасная рабочая зона",
        "Памятка для смены ООО \"ДетаЛит\"",
        [
            ("Средства защиты", "Очки, перчатки, спецодежда и обувь должны быть надеты до входа в производственную зону."),
            ("Проходы и зоны", "Не загромождайте проходы, держите доступ к аварийной кнопке, аптечке и огнетушителю."),
            ("Сигналы отклонений", "Искры, запах гари, шум, вибрация или перегрев — повод остановить операцию и позвать мастера."),
            ("Фиксация события", "Любое происшествие записывается в журнал смены с временем, участком и ответственным."),
        ],
        "#1f4e79",
    )
    make_png(
        "process-flow.png",
        "Маршрут изготовления отливки",
        "От задания смены до передачи в ОТК",
        [
            ("Подготовка формы", "Очистка, смазка, фиксация формы и проверка маркировки партии."),
            ("Плавка и заливка", "Контроль температуры, состава и стабильности заливки без рывков."),
            ("Охлаждение", "Выдержка по регламенту без досрочного вскрытия формы."),
            ("Контроль качества", "Визуальный контроль, измерение критических размеров, запись результата."),
        ],
        "#0f766e",
    )
    make_png(
        "quality-checkpoints.png",
        "Контрольные точки качества",
        "Что сотрудник проверяет на каждом изделии",
        [
            ("Поверхность", "Ищем трещины, раковины, непроливы, подгар и следы механических повреждений."),
            ("Геометрия", "Сверяем базовые размеры, плоскостность, отверстия и посадочные поверхности."),
            ("Маркировка", "Партия, смена и маршрутная карта должны совпадать."),
            ("Решение", "Годное изделие передается дальше, спорное — в изолятор брака и к мастеру."),
        ],
        "#b45309",
    )
    make_docx()
    make_pdf()
    print(f"Generated course assets in {OUT}")


if __name__ == "__main__":
    main()
